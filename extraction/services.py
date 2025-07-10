# # extraction/services.py - Version corrigée et robuste


# extraction/services.py - Version optimisée avec recherche EMA robuste
import os
import re
import json
import requests
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional
import logging

import spacy
import docx
from PyPDF2 import PdfReader
from django.utils import timezone
from langdetect import detect
from dotenv import load_dotenv

# Charger les variables d'environnement
load_dotenv()

logger = logging.getLogger(__name__)

# Chargement conditionnel des modèles spaCy
try:
    NLP_FR = spacy.load("fr_core_news_sm")
    NLP_EN = spacy.load("en_core_web_sm")
    SPACY_AVAILABLE = True
except OSError:
    SPACY_AVAILABLE = False
    logger.warning("⚠️ Modèles spaCy non disponibles")

# Stopwords pour nettoyage
STOPWORDS = {
    "le", "la", "les", "de", "des", "du", "un", "une", "et", "en", "à", "dans", "que", "qui",
    "pour", "par", "sur", "avec", "au", "aux", "ce", "ces", "se", "ses", "est",
    "the", "and", "of", "to", "in", "that", "it", "is", "was", "for", "on", "are", "with",
    "as", "i", "at", "be", "by", "this"
}


class RobustTitleExtractor:
    """Extracteur de titre robuste avec scores de confiance"""

    def __init__(self):
        self.spacy_available = SPACY_AVAILABLE

    def clean_text(self, text: str) -> str:
        """Nettoyage du texte pour l'analyse"""
        text = re.sub(r"[^0-9A-Za-zÀ-ÖØ-öø-ÿ\s\.,;:\-'\(\)]", " ", text)
        text = re.sub(r"\s+", " ", text).strip()
        words = text.split()
        return " ".join(w for w in words if w.lower() not in STOPWORDS)

    def extract_title(self, text: str, document_title: str = None) -> Dict:
        """Extraction de titre avec scores de confiance"""
        if not text or len(text.strip()) < 20:
            return self._fallback_result(document_title)

        cleaned_text = self.clean_text(text)
        lines = cleaned_text.split('\n')

        # Chercher dans les premières lignes
        title_candidates = self._find_title_candidates(lines[:20])

        if not title_candidates:
            title_candidates = self._find_title_candidates(lines[20:70])

        best_title = self._select_best_title(title_candidates, document_title)

        return {
            'title': best_title['title'],
            'confidence': best_title['confidence'],
            'method': best_title['method'],
            'position': best_title.get('position', 0),
            'alternatives': [c for c in title_candidates if c['title'] != best_title['title']]
        }

    def _find_title_candidates(self, lines: List[str]) -> List[Dict]:
        """Trouver des candidats titres dans les lignes"""
        candidates = []

        for i, line in enumerate(lines):
            line = line.strip()
            if not line or len(line) < 10 or len(line) > 200:
                continue

            confidence = self._calculate_title_confidence(line, i)
            if confidence < 0.5:
                continue

            candidates.append({
                'title': line,
                'confidence': confidence,
                'method': 'content_analysis',
                'position': i
            })

        return candidates

    def _calculate_title_confidence(self, line: str, position: int) -> float:
        """Calculer un score de confiance pour une ligne candidate"""
        confidence = 0.0

        # Bonus pour les premières lignes
        if position < 5:
            confidence += 0.3

        # Bonus pour les titres en majuscules ou title case
        if line.istitle() or line.isupper():
            confidence += 0.2

        # Bonus pour la longueur modérée
        if 20 <= len(line) <= 120:
            confidence += 0.2

        # Pénalité pour les métadonnées
        metadata_terms = ['page', 'version', 'copyright', 'email:', 'http', '@']
        if any(term in line.lower() for term in metadata_terms):
            confidence -= 0.5

        # Pénalité pour trop de caractères spéciaux
        special_count = len(re.findall(r'[^\w\s\-\'\",\.\(\)]', line))
        if special_count > len(line) * 0.2:
            confidence -= 0.3

        return max(0.0, min(1.0, confidence))

    def _select_best_title(self, candidates: List[Dict], document_title: str = None) -> Dict:
        """Sélectionner le meilleur titre parmi les candidats"""
        if not candidates:
            return self._fallback_result(document_title)

        sorted_candidates = sorted(candidates, key=lambda x: x['confidence'], reverse=True)

        # Vérifier si le titre du document est valide
        if document_title and self._is_document_title_valid(document_title):
            doc_title_confidence = self._calculate_title_confidence(document_title, 0)
            if doc_title_confidence > sorted_candidates[0]['confidence']:
                return {
                    'title': document_title,
                    'confidence': doc_title_confidence,
                    'method': 'document_title'
                }

        return sorted_candidates[0]

    def _is_document_title_valid(self, title: str) -> bool:
        """Vérifier si le titre du document est valide"""
        if not title or len(title.strip()) < 5:
            return False

        title_lower = title.lower()
        generic_terms = ['document', 'untitled', 'file', 'test', 'temp', 'sans titre']
        if any(term in title_lower for term in generic_terms):
            return False

        if title_lower.endswith(('.pdf', '.doc', '.docx', '.txt')):
            return False

        return True

    def _fallback_result(self, document_title: str = None) -> Dict:
        """Résultat de fallback"""
        title = document_title if document_title else "Document sans titre"
        return {
            'title': title,
            'confidence': 0.1,
            'method': 'fallback',
            'position': 0,
            'alternatives': []
        }


class RobustMetadataExtractor:
    """Extracteur de métadonnées robuste avec patterns"""

    def __init__(self):
        self.title_extractor = RobustTitleExtractor()

    def extract_robust_patterns(self, text: str, file_type: str, document_title: str = None) -> Dict:
        """Extraction robuste par patterns pour tous types de documents"""

        metadata = {
            "title": "",
            "document_type": "document",  # Changé de "type" à "document_type"
            "context": "general",
            "language": "fr",
            "publication_date": None,
            "source": "",
            "version": "",
            "country": "",
            "url_source": "",
            "confidence_scores": {}
        }

        try:
            # 1. Extraction du titre
            title_result = self.title_extractor.extract_title(text, document_title)
            metadata['title'] = title_result['title']
            metadata['confidence_scores']['title'] = title_result['confidence']

            # 2. Détection de langue
            lang_result = self._detect_language_robust(text)
            metadata['language'] = lang_result['language']
            metadata['confidence_scores']['language'] = lang_result['confidence']

            # 3. Type de document
            type_result = self._detect_document_type_robust(text)
            metadata['document_type'] = type_result['type']  # Mapper vers document_type
            metadata['confidence_scores']['document_type'] = type_result['confidence']

            # 4. Contexte
            context_result = self._detect_context_robust(text)
            metadata['context'] = context_result['context']
            metadata['confidence_scores']['context'] = context_result['confidence']

            # 5. Date de publication
            date_result = self._extract_date_robust(text)
            if date_result:
                metadata['publication_date'] = date_result['date']
                metadata['confidence_scores']['publication_date'] = date_result['confidence']
            else:
                metadata['confidence_scores']['publication_date'] = 0.0

            # 6. Source
            source_result = self._detect_source_robust(text)
            metadata['source'] = source_result['source']
            metadata['confidence_scores']['source'] = source_result['confidence']

            # 7. Version
            version_result = self._extract_version_robust(text)
            metadata['version'] = version_result['version']
            metadata['confidence_scores']['version'] = version_result['confidence']

            # 8. Pays
            country_result = self._detect_country_robust(text)
            metadata['country'] = country_result['country']
            metadata['confidence_scores']['country'] = country_result['confidence']

        except Exception as e:
            logger.error(f"Erreur extraction robuste: {e}")
            metadata['confidence_scores'] = {k: 0.1 for k in metadata}

        return metadata

    def _detect_language_robust(self, text: str) -> Dict:
        """Détection de langue robuste"""
        text_lower = text.lower()

        # Mots indicateurs français
        french_words = ['le', 'la', 'les', 'de', 'du', 'des', 'et', 'à', 'dans', 'pour', 'sur', 'avec', 'par']
        french_count = sum(1 for word in french_words if f' {word} ' in text_lower)

        # Mots indicateurs anglais
        english_words = ['the', 'and', 'of', 'to', 'in', 'for', 'on', 'with', 'by', 'from']
        english_count = sum(1 for word in english_words if f' {word} ' in text_lower)

        if french_count > english_count * 1.2:
            return {"language": "fr", "confidence": min(0.7 + french_count * 0.01, 0.95)}
        elif english_count > french_count * 1.2:
            return {"language": "en", "confidence": min(0.7 + english_count * 0.01, 0.95)}
        else:
            return {"language": "fr", "confidence": 0.6}

    def _detect_document_type_robust(self, text: str) -> Dict:
        """Détection du type de document"""
        text_lower = text.lower()

        types = {
            "guideline": ["guideline", "guidance", "guide", "instruction", "directive"],
            "regulation": ["regulation", "règlement", "law", "loi", "directive", "decree"],
            "report": ["rapport", "report", "analysis", "analyse", "study", "étude"],
            "procedure": ["procedure", "procédure", "method", "méthode", "protocol", "protocole"],
            "standard": ["standard", "norme", "specification", "requirement"],
            "directive": ["directive", "instruction", "consigne"]
        }

        max_score = 0
        best_type = "other"

        for doc_type, keywords in types.items():
            score = sum(text_lower.count(keyword) for keyword in keywords)
            if score > max_score:
                max_score = score
                best_type = doc_type

        confidence = min(0.5 + max_score * 0.1, 0.9) if max_score > 0 else 0.3
        return {"type": best_type, "confidence": confidence}

    def _detect_context_robust(self, text: str) -> Dict:
        """Détection du contexte"""
        text_lower = text.lower()

        contexts = {
            "pharmaceutical": ["pharmaceutical", "medicine", "drug", "medicament", "thérapeutique"],
            "medical": ["medical", "médical", "clinical", "clinique", "patient", "treatment"],
            "regulatory": ["regulation", "règlement", "compliance", "authorisation", "approval"],
            "legal": ["legal", "juridique", "law", "loi", "court", "tribunal"]
        }

        max_score = 0
        best_context = "general"

        for context, keywords in contexts.items():
            score = sum(text_lower.count(keyword) for keyword in keywords)
            if score > max_score:
                max_score = score
                best_context = context

        confidence = min(0.5 + max_score * 0.1, 0.9) if max_score > 0 else 0.3
        return {"context": best_context, "confidence": confidence}

    def _extract_date_robust(self, text: str) -> Optional[Dict]:
        """Extraction de date robuste"""
        date_patterns = [
            r'\b(\d{1,2})[/-](\d{1,2})[/-](\d{4})\b',  # DD/MM/YYYY
            r'\b(\d{4})[/-](\d{1,2})[/-](\d{1,2})\b',  # YYYY/MM/DD
            r'\b(\d{1,2})\s+(janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre)\s+(\d{4})\b',
            r'\b(\d{1,2})\s+(january|february|march|april|may|june|july|august|september|october|november|december)\s+(\d{4})\b',
        ]

        for pattern in date_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                date_str = match.group()
                parsed_date = self._parse_date_string(date_str)
                if parsed_date:
                    return {"date": parsed_date, "confidence": 0.7}

        return None

    def _parse_date_string(self, date_str: str) -> Optional[str]:
        """Parser une date string avec gestion robuste des formats"""
        try:
            # Essayer avec dateparser en premier si disponible
            try:
                import dateparser
                parsed = dateparser.parse(date_str, languages=['fr', 'en'])
                if parsed:
                    return parsed.strftime('%d %B %Y')
            except ImportError:
                pass  # dateparser non disponible, continuer avec les formats manuels

            # Fallback manuel avec formats courants
            formats = [
                '%d/%m/%Y', '%d-%m-%Y', '%Y/%m/%d', '%Y-%m-%d',
                '%d %B %Y', '%d %b %Y', '%B %d, %Y', '%b %d, %Y',
                '%d.%m.%Y', '%Y.%m.%d'
            ]

            for fmt in formats:
                try:
                    parsed = datetime.strptime(date_str, fmt)
                    return parsed.strftime('%d %B %Y')
                except ValueError:
                    continue

            # Si aucun format ne marche, essayer d'extraire l'année
            import re
            year_match = re.search(r'\b(19|20)\d{2}\b', date_str)
            if year_match:
                year = int(year_match.group())
                if 1990 <= year <= 2030:
                    return f"01 January {year}"

        except Exception as e:
            logger.warning(f"Erreur parsing date '{date_str}': {e}")

        return None

    def _detect_source_robust(self, text: str) -> Dict:
        """Détection de source robuste"""
        text_lower = text.lower()

        sources = {
            "EMA": ["ema", "european medicines agency", "agence européenne des médicaments"],
            "FDA": ["fda", "food and drug administration"],
            "ANSM": ["ansm", "agence nationale de sécurité du médicament"],
            "WHO": ["who", "world health organization", "organisation mondiale de la santé"]
        }

        for source, keywords in sources.items():
            if any(keyword in text_lower for keyword in keywords):
                return {"source": source, "confidence": 0.8}

        return {"source": "", "confidence": 0.0}

    def _extract_version_robust(self, text: str) -> Dict:
        """Extraction de version robuste"""
        version_patterns = [
            r'version\s*[:=]?\s*([0-9]+(?:\.[0-9]+)*)',
            r'v\.?\s*([0-9]+(?:\.[0-9]+)*)',
            r'révision\s*[:=]?\s*([0-9]+)',
            r'rev\.?\s*([0-9]+)'
        ]

        for pattern in version_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return {"version": match.group(1), "confidence": 0.8}

        return {"version": "", "confidence": 0.0}

    def _detect_country_robust(self, text: str) -> Dict:
        """Détection de pays robuste"""
        text_lower = text.lower()

        countries = {
            "EU": ["european", "europe", "eu", "européen", "europa"],
            "US": ["united states", "usa", "america", "états-unis"],
            "FR": ["france", "français", "french"],
            "UK": ["united kingdom", "britain", "royaume-uni"]
        }

        for country, keywords in countries.items():
            if any(keyword in text_lower for keyword in keywords):
                return {"country": country, "confidence": 0.8}

        return {"country": "", "confidence": 0.0}


class DocumentTextExtractor:
    """Extracteur de texte avec gestion d'erreurs améliorée"""

    @staticmethod
    def extract_text_from_file(file_path: str, file_type: str) -> str:
        """Extraction avec nettoyage du texte"""
        if not os.path.exists(file_path):
            logger.error(f"Fichier non trouvé: {file_path}")
            return ""

        try:
            if file_type.lower() == 'pdf':
                text = DocumentTextExtractor._extract_from_pdf(file_path)
            elif file_type.lower() in ['docx', 'doc']:
                text = DocumentTextExtractor._extract_from_docx(file_path)
            elif file_type.lower() == 'txt':
                text = DocumentTextExtractor._extract_from_txt(file_path)
            else:
                logger.warning(f"Type non supporté: {file_type}")
                return ""

            # Nettoyer le texte
            text = re.sub(r'\s+', ' ', text).strip()
            return text

        except Exception as e:
            logger.error(f"Erreur extraction texte: {e}")
            return ""

    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """Extraction PDF avec gestion d'erreurs"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                for i, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Erreur extraction page {i}: {e}")
                        continue
        except Exception as e:
            logger.error(f"Erreur PDF: {e}")
            return ""
        return text

    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """Extraction DOCX avec gestion d'erreurs"""
        try:
            doc = docx.Document(file_path)
            text_parts = []

            # Extraire les paragraphes
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extraire les tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Erreur DOCX: {e}")
            return ""

    @staticmethod
    def _extract_from_txt(file_path: str) -> str:
        """Extraction TXT avec détection d'encodage"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Erreur lecture fichier avec {encoding}: {e}")
                continue

        logger.error("Impossible de décoder le fichier texte")
        return ""


class MistralAIService:
    """Service d'extraction de métadonnées robuste avec fallback amélioré"""

    def __init__(self):
        self.title_extractor = RobustTitleExtractor()
        self.metadata_extractor = RobustMetadataExtractor()

        # Vérifier la disponibilité d'Ollama/API externes
        self.ollama_available = self._check_ollama_availability()

    def _check_ollama_availability(self) -> bool:
        """Vérifier si l'API Mistral ou Ollama est disponible"""
        # Vérifier l'API Mistral
        api_key = os.getenv("MISTRAL_API_KEY")
        if api_key:
            return True

        # Vérifier Ollama local
        try:
            import requests
            response = requests.get("http://localhost:11434/api/version", timeout=5)
            return response.status_code == 200
        except:
            return False

    def extract_metadata_with_confidence(self, text: str, file_type: str, document_title: str = None,
                                         source_url: str = "") -> Dict:
        """Extraction avec scores de confiance - NE PAS utiliser extract_metadata_with_ema ici"""
        if not text or len(text.strip()) < 10:
            return self._get_minimal_metadata(document_title, source_url)

        # Extraction du titre avec confiance
        title_result = self.title_extractor.extract_title(text, document_title)

        # Extraction des autres métadonnées
        try:
            metadata = self._extract_with_mistral_api(text, file_type, document_title, source_url)

            # IMPORTANT : Garder le titre Mistral s'il est meilleur
            mistral_title = metadata.get('title', '').strip()
            if mistral_title and mistral_title.lower() not in ['document sans titre', 'untitled', '']:
                logger.info(f"✅ Titre Mistral conservé: '{mistral_title}'")
                # Ne pas écraser avec le titre du RobustTitleExtractor
            else:
                # Utiliser le titre du RobustTitleExtractor si Mistral n'a pas trouvé
                metadata['title'] = title_result['title']
                metadata['confidence_scores']['title'] = title_result['confidence']

        except Exception as e:
            logger.warning(f"⚠️ Échec Mistral, fallback vers extraction robuste: {e}")
            metadata = self.metadata_extractor.extract_robust_patterns(text, file_type, document_title)

        # Ajouter URL source
        metadata['url_source'] = source_url

        # NE PAS faire de recherche EMA ici, elle sera faite séparément
        # metadata['ema_data'] = self._get_empty_ema_data()

        # Calcul du score global
        metadata['quality_score'] = self._calculate_quality_score(metadata.get('confidence_scores', {}))

        return metadata

    def extract_metadata_with_ema(self, text: str, file_type: str, document_title: str,
                                  ema_page_url: str = None) -> Dict:
        """Extraction des métadonnées avec données EMA - VERSION AVEC LOGS"""

        logger.info(f"🔍 MistralAI extraction pour: '{document_title[:50]}...'")
        logger.info(f"🔗 URL EMA fournie: {ema_page_url}")
        base_metadata = self.extract_metadata_with_confidence(text, file_type, document_title)
        print("📊 Métadonnées de base extraites:", base_metadata)
        try:
            # Extraction métadonnées de base
            base_metadata = self.extract_metadata_with_confidence(text, file_type, document_title)
            print("📊 Métadonnées de base extraites:", base_metadata)
            logger.info(f"📊 Métadonnées de base extraites: {base_metadata.get('title', 'N/A')[:50]}...")

            # Recherche EMA si URL fournie
            ema_data = {}
            if ema_page_url:
                logger.info(f"🔍 Recherche EMA FRAICHE pour: {ema_page_url}")

                from .url_services import EMAPageExtractor
                ema_extractor = EMAPageExtractor()

                extracted_title = base_metadata.get('title', document_title)
                logger.info(f"🎯 Recherche EMA avec titre: '{extracted_title}'")

                ema_result = ema_extractor.extract_metadata_from_ema_page(ema_page_url, extracted_title)

                if ema_result.get('success'):
                    ema_data = {
                        'original_publication_date': ema_result.get('original_publication_date'),
                        'ema_publication_date': ema_result.get('ema_publication_date'),
                        'ema_source_url': ema_result.get('ema_source_url', ''),
                        'ema_title': ema_result.get('ema_title', ''),
                        'ema_reference': ema_result.get('ema_reference', ''),
                        'similarity_score': ema_result.get('similarity_score', 0.0),
                        'search_performed': True,
                        'search_timestamp': ema_result.get('search_timestamp'),
                        'success': True
                    }
                    logger.info(f"✅ NOUVELLES données EMA extraites:")
                    logger.info(f"  - ema_title: '{ema_data.get('ema_title', '')}'")
                    logger.info(f"  - ema_source_url: '{ema_data.get('ema_source_url', '')}'")
                    logger.info(f"  - original_publication_date: {ema_data.get('original_publication_date')}")
                else:
                    logger.warning(f"⚠️ Échec extraction EMA: {ema_result.get('error')}")
                    ema_data = {
                        'search_performed': True,
                        'search_timestamp': timezone.now().isoformat(),
                        'error': ema_result.get('error'),
                        'success': False
                    }
            else:
                logger.info(f"📝 Pas d'URL EMA fournie, pas de recherche EMA")
                ema_data = {
                    'search_performed': False,
                    'success': False
                }

            # Combiner les données
            base_metadata['ema_data'] = ema_data
            print(f"📊 Métadonnées combinées: {base_metadata}")

            logger.info(f"📊 Métadonnées finales:")
            logger.info(f"  - Titre: '{base_metadata.get('title', 'N/A')[:50]}...'")
            logger.info(f"  - EMA data incluses: {ema_data.get('success', False)}")

            return base_metadata

        except Exception as e:
            logger.error(f"❌ Erreur extraction avec EMA: {e}")
            # Retourner au moins les métadonnées de base
            try:
                base_metadata = self.extract_metadata_with_confidence(text, file_type, document_title)
            except:
                base_metadata = {'title': document_title}

            base_metadata['ema_data'] = {
                'search_performed': False,
                'error': str(e),
                'success': False
            }
            return base_metadata

    def _extract_with_mistral_api(self, text: str, file_type: str, document_title: str, source_url: str) -> Dict:
        """Extraction avec l'API Mistral - PRÉSERVATION TITRE"""
        result = call_mistral_with_confidence(text, source_url)

        if not result or 'metadata' not in result:
            raise Exception("Réponse Mistral invalide")

        metadata = result['metadata']

        # CORRECTION CRITIQUE : Vérifier le titre extrait
        extracted_title = metadata.get('title', '').strip()
        logger.info(f"🔍 Titre dans réponse Mistral: '{extracted_title}'")

        if not extracted_title or extracted_title.lower() in ['document sans titre', 'untitled', '', 'none']:
            logger.warning(f"⚠️ Mistral n'a pas trouvé de titre valide: '{extracted_title}'")
            if document_title and document_title.lower() not in ['document sans titre', 'untitled', '']:
                metadata['title'] = document_title
                logger.info(f"🔄 Utilisation titre document fourni: '{document_title}'")
            else:
                metadata['title'] = "Titre non détecté par LLM"
                logger.warning("❌ Aucun titre de fallback disponible")
        else:
            logger.info(f"✅ Titre Mistral valide: '{extracted_title}'")

        metadata['confidence_scores'] = result.get('confidence_scores', {})
        metadata['extraction_reasoning'] = result.get('extraction_reasoning', {})

        return metadata

    def _extract_with_mistral_api(self, text: str, file_type: str, document_title: str, source_url: str) -> Dict:
        """Extraction avec l'API Mistral"""
        result = call_mistral_with_confidence(text, source_url)

        if not result or 'metadata' not in result:
            raise Exception("Réponse Mistral invalide")

        metadata = result['metadata']

        # CORRECTION : Vérifier que le titre est bien extrait
        extracted_title = metadata.get('title', '').strip()
        logger.info(f"🔍 Titre extrait par Mistral: '{extracted_title}'")

        if not extracted_title or extracted_title.lower() in ['document sans titre', 'untitled', '']:
            logger.warning("⚠️ Mistral n'a pas trouvé de titre valide")
            if document_title:
                metadata['title'] = document_title
                logger.info(f"📝 Utilisation titre document fourni: '{document_title}'")

        metadata['confidence_scores'] = result.get('confidence_scores', {})
        metadata['extraction_reasoning'] = result.get('extraction_reasoning', {})

        return metadata

    def _extract_with_mistral_api(self, text: str, file_type: str, document_title: str, source_url: str) -> Dict:
        """Extraction avec l'API Mistral"""
        result = call_mistral_with_confidence(text, source_url)

        if not result or 'metadata' not in result:
            raise Exception("Réponse Mistral invalide")

        metadata = result['metadata']
        metadata['confidence_scores'] = result.get('confidence_scores', {})
        metadata['extraction_reasoning'] = result.get('extraction_reasoning', {})

        return metadata

    def _get_best_title_for_ema(self, document_title: str, extracted_title: str) -> str:
        """Déterminer le meilleur titre pour la recherche EMA"""
        if extracted_title and len(extracted_title.strip()) > 10:
            return extracted_title
        if document_title and len(document_title.strip()) > 10:
            return document_title
        return extracted_title or document_title or ""

    def _perform_ema_search(self, title: str, ema_page_url: str = None) -> Dict:
        """Effectuer la recherche EMA ou extraction directe si URL fournie"""
        try:
            # Si on a une URL EMA directe, l'utiliser
            if ema_page_url:
                logger.info(f"🔗 URL EMA fournie, extraction directe: {ema_page_url}")
                return fetch_ema_data_from_url(ema_page_url)

            # Sinon, faire une recherche normale
            logger.info(f"🔍 Recherche EMA pour: '{title}'")

            if not title or title.lower() in ['document sans titre', 'untitled', '']:
                logger.warning("⚠️ Titre invalide pour recherche EMA")
                return self._get_empty_ema_data(title)

            ema_raw_data = fetch_ema_dates_by_title(title)

            if ema_raw_data and ema_raw_data.get('search_successful'):
                logger.info(f"✅ Données EMA trouvées: {ema_raw_data.get('ema_title', '')}")

                mapped_data = {
                    'ema_title': ema_raw_data.get('ema_title', ''),
                    'ema_source_url': ema_raw_data.get('ema_source_url', ''),
                    'original_publication_date': ema_raw_data.get('original_publication_date'),
                    'ema_publication_date': ema_raw_data.get('ema_publication_date'),
                    'ema_reference': ema_raw_data.get('ema_reference', ''),
                    'similarity_score': ema_raw_data.get('similarity_score', 0.0),
                    'search_successful': True,
                    'search_performed': True,
                    'search_title_used': ema_raw_data.get('search_term_used', title),
                    'search_timestamp': ema_raw_data.get('search_timestamp', datetime.now().isoformat()),
                }

                return mapped_data
            else:
                logger.warning(f"❌ Aucun résultat EMA pour: '{title}'")
                return self._get_empty_ema_data(title)

        except Exception as e:
            logger.error(f"❌ Erreur recherche EMA: {e}")
            return self._get_empty_ema_data(title, str(e))

    def _get_empty_ema_data(self, title: str = "", error: str = None) -> Dict:
        """Données EMA vides avec la structure attendue par tasks.py"""
        return {
            'ema_title': '',
            'ema_source_url': '',
            'original_publication_date': None,
            'ema_publication_date': None,
            'ema_reference': '',
            'similarity_score': 0.0,
            'search_successful': False,
            'search_performed': True,
            'search_title_used': title,
            'search_timestamp': datetime.now().isoformat(),
            'search_error': error,
        }

    def _calculate_quality_score(self, confidence_scores: Dict) -> float:
        """Calculer un score de qualité global"""
        if not confidence_scores:
            return 0.0

        weights = {
            'title': 1.5,
            'document_type': 1.2,  # Changé de 'type' à 'document_type'
            'language': 1.0,
            'publication_date': 1.3,
            'source': 1.2,
            'context': 1.0,
            'version': 0.7,
            'country': 0.8
        }

        total = 0.0
        weight_sum = 0.0

        for field, weight in weights.items():
            if field in confidence_scores:
                total += confidence_scores[field] * weight
                weight_sum += weight

        return total / weight_sum if weight_sum > 0 else 0.0

    def _get_minimal_metadata(self, document_title: str = None, source_url: str = "") -> Dict:
        """Métadonnées minimales"""
        return {
            "title": document_title or "Document sans titre",
            "document_type": "other",  # Changé de "type" à "document_type"
            "context": "general",
            "language": "fr",
            "publication_date": None,
            "source": "",
            "version": "",
            "country": "",
            "url_source": source_url,
            "ema_data": self._get_empty_ema_data(),
            "confidence_scores": {
                "title": 0.1,
                "document_type": 0.1,  # Changé de "type" à "document_type"
                "context": 0.1,
                "language": 0.5,
                "publication_date": 0.0,
                "source": 0.0,
                "version": 0.0,
                "country": 0.0
            },
            "quality_score": 0.1
        }


class NLPAnnotationService:
    """Service d'annotation automatique simplifié"""

    def __init__(self):
        self.spacy_available = SPACY_AVAILABLE

    def auto_annotate_document(self, text: str) -> List[Dict]:
        """Annotation automatique avec patterns robustes"""
        annotations = []

        # Patterns généraux améliorés
        patterns = {
            'ORGANISATION': [
                r'\b(EMA|FDA|European Medicines Agency|Agence européenne des médicaments)\b',
                r'\b(ANSM|Agence nationale de sécurité du médicament)\b',
                r'\b(Commission européenne|European Commission)\b',
                r'\b(WHO|World Health Organization|Organisation mondiale de la santé)\b'
            ],
            'PROCEDURE': [
                r'\b(ITF|Briefing Meeting|Application|Demande)\b',
                r'\b(Autorisation|Authorization|Approval)\b',
                r'\b(Procédure|Procedure|Protocol)\b'
            ],
            'INSTRUCTION': [
                r'\b(Instructions|Guidelines|Directives|Procédures)\b',
                r'\b(Guide|Manuel|Documentation)\b'
            ],
            'DATE': [
                r'\b\d{1,2}[/-]\d{1,2}[/-]\d{4}\b',
                r'\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b'
            ],
            'REFERENCE': [
                r'\b[A-Z]{2,}/[A-Z]{2,}/\d+\b',
                r'\bRef\.\s*[A-Z0-9/-]+\b'
            ],
            'MEDICAL_TERM': [
                r'\b(pharmaceutical|medicine|drug|medicament|thérapeutique)\b',
                r'\b(clinical|clinique|patient|treatment|traitement)\b'
            ]
        }

        for entity_type, pattern_list in patterns.items():
            for pattern in pattern_list:
                for match in re.finditer(pattern, text, re.IGNORECASE):
                    annotations.append({
                        'text': match.group(),
                        'start_position': match.start(),
                        'end_position': match.end(),
                        'entity_type': entity_type,
                        'confidence_score': 0.7,
                        'is_automatic': True,
                        'source': 'pattern_enhanced'
                    })

        return annotations

    def extract_entities_with_spacy(self, text: str) -> List[Dict]:
        """Extraction avec spaCy si disponible"""
        if not self.spacy_available:
            return self.auto_annotate_document(text)

        try:
            # Détecter la langue et utiliser le bon modèle
            lang = detect(text[:1000]) if len(text) > 1000 else detect(text)
            nlp = NLP_FR if lang == 'fr' else NLP_EN

            doc = nlp(text[:1000000])  # Limite pour éviter les timeouts

            entities = []
            for ent in doc.ents:
                entities.append({
                    'text': ent.text,
                    'start_position': ent.start_char,
                    'end_position': ent.end_char,
                    'entity_type': ent.label_,
                    'confidence_score': 0.8,
                    'is_automatic': True,
                    'source': 'spacy'
                })

            return entities

        except Exception as e:
            logger.warning(f"Erreur spaCy: {e}, fallback vers patterns")
            return self.auto_annotate_document(text)


import requests
from bs4 import BeautifulSoup
from urllib.parse import quote
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

import requests
from bs4 import BeautifulSoup
from urllib.parse import quote, urljoin
from datetime import datetime
import logging
import re
from difflib import SequenceMatcher
from typing import Dict, List, Optional, Tuple
import time

logger = logging.getLogger(__name__)

BASE_URL = "https://www.ema.europa.eu"
SEARCH_ENDPOINT = f"{BASE_URL}/en/search"


class EMASearchEngine:
    """Moteur de recherche robuste pour l'EMA avec techniques avancées"""

    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.9',
            'Accept-Encoding': 'gzip, deflate, br',
            'Connection': 'keep-alive',
            'Cache-Control': 'no-cache'
        })

    def normalize_title(self, title: str) -> str:
        """Normalise un titre pour la comparaison"""
        if not title:
            return ""

        # Convertir en minuscules
        normalized = title.lower()

        # Remplacer les caractères spéciaux par des espaces
        normalized = re.sub(r'[^\w\s]', ' ', normalized)

        # Supprimer les mots vides courants
        stopwords = {
            'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'from',
            'a', 'an', 'as', 'how', 'what', 'where', 'when', 'why', 'which', 'who', 'that', 'this',
            'le', 'la', 'les', 'de', 'du', 'des', 'et', 'à', 'dans', 'pour', 'sur', 'avec', 'par'
        }

        words = normalized.split()
        words = [w for w in words if w not in stopwords and len(w) > 2]

        return ' '.join(words)

    def extract_keywords(self, title: str, max_keywords: int = 5) -> List[str]:
        """Extrait les mots-clés les plus importants d'un titre"""
        normalized = self.normalize_title(title)
        words = normalized.split()

        # Prioriser les mots longs et spécialisés
        word_scores = {}
        for word in words:
            score = len(word)  # Les mots longs sont plus spécifiques

            # Bonus pour les termes médicaux/pharmaceutiques
            medical_indicators = ['guidance', 'guideline', 'medicine', 'drug', 'therapeutic',
                                  'clinical', 'treatment', 'procedure', 'protocol', 'meeting',
                                  'briefing', 'application', 'authorization', 'approval']

            if any(indicator in word for indicator in medical_indicators):
                score += 5

            word_scores[word] = score

        # Retourner les mots avec les meilleurs scores
        sorted_words = sorted(word_scores.items(), key=lambda x: x[1], reverse=True)
        return [word for word, score in sorted_words[:max_keywords]]

    def calculate_similarity(self, text1: str, text2: str) -> float:
        """Calcule la similarité entre deux textes"""
        if not text1 or not text2:
            return 0.0

        norm1 = self.normalize_title(text1)
        norm2 = self.normalize_title(text2)

        # Similarité basée sur la séquence
        seq_similarity = SequenceMatcher(None, norm1, norm2).ratio()

        # Similarité basée sur les mots communs
        words1 = set(norm1.split())
        words2 = set(norm2.split())

        if not words1 and not words2:
            return 1.0
        elif not words1 or not words2:
            return 0.0

        intersection = words1.intersection(words2)
        union = words1.union(words2)
        word_similarity = len(intersection) / len(union)

        # Moyenne pondérée
        return (seq_similarity * 0.4) + (word_similarity * 0.6)

    def generate_search_strategies(self, title: str) -> List[Dict[str, str]]:
        """Génère différentes stratégies de recherche"""
        strategies = []

        # 1. Recherche exacte avec guillemets
        strategies.append({
            'query': f'"{title}"',
            'type': 'exact_quoted',
            'description': 'Recherche exacte avec guillemets'
        })

        # 2. Recherche exacte sans guillemets
        strategies.append({
            'query': title,
            'type': 'exact_unquoted',
            'description': 'Recherche exacte sans guillemets'
        })

        # 3. Mots-clés principaux
        keywords = self.extract_keywords(title, 4)
        if keywords:
            strategies.append({
                'query': ' '.join(keywords),
                'type': 'keywords',
                'description': f'Mots-clés principaux: {", ".join(keywords)}'
            })

        # 4. Première moitié du titre + mots-clés
        words = title.split()
        if len(words) > 3:
            first_half = ' '.join(words[:len(words) // 2])
            key_terms = self.extract_keywords(title, 2)
            combined = f"{first_half} {' '.join(key_terms)}"
            strategies.append({
                'query': combined,
                'type': 'partial_keywords',
                'description': 'Première partie + mots-clés'
            })

        # 5. Termes spécialisés seulement
        specialized_terms = []
        specialized_patterns = [
            r'\b(guidance|guideline|procedure|protocol|meeting|briefing)\b',
            r'\b(application|authorization|approval|submission)\b',
            r'\b(ITF|EMA|FDA|CHMP|PDCO|CAT)\b',
            r'\b[A-Z]{2,}\b'  # Acronymes
        ]

        for pattern in specialized_patterns:
            matches = re.findall(pattern, title, re.IGNORECASE)
            specialized_terms.extend(matches)

        if specialized_terms:
            strategies.append({
                'query': ' '.join(set(specialized_terms)),
                'type': 'specialized',
                'description': f'Termes spécialisés: {", ".join(set(specialized_terms))}'
            })

        # 6. Recherche par acronymes uniquement
        acronyms = re.findall(r'\b[A-Z]{2,}\b', title)
        if acronyms:
            strategies.append({
                'query': ' '.join(acronyms),
                'type': 'acronyms',
                'description': f'Acronymes: {", ".join(acronyms)}'
            })

        return strategies

    def search_ema(self, query: str, max_results: int = 10) -> List[BeautifulSoup]:
        """Effectue une recherche sur le site EMA"""
        try:
            params = {
                'search_api_fulltext': query,
                'f[0]': 'ema_search_entity_is_document:Document'
            }

            logger.debug(f"🔍 Recherche EMA avec query: '{query}'")

            response = self.session.get(SEARCH_ENDPOINT, params=params, timeout=30)
            response.raise_for_status()

            soup = BeautifulSoup(response.text, 'html.parser')

            # Chercher les résultats avec plusieurs sélecteurs possibles
            result_selectors = [
                '.search-results .search-result',
                '.views-row',
                '.item-list .item',
                '[class*="search-result"]',
                '.bcl-listing .row > div'
            ]

            results = []
            for selector in result_selectors:
                found_results = soup.select(selector)
                if found_results:
                    logger.debug(f"✅ Trouvé {len(found_results)} résultats avec sélecteur: {selector}")
                    results = found_results[:max_results]
                    break

            if not results:
                logger.warning(f"❌ Aucun résultat trouvé pour: '{query}'")

            return results

        except requests.RequestException as e:
            logger.error(f"❌ Erreur réseau lors de la recherche: {e}")
            return []
        except Exception as e:
            logger.error(f"❌ Erreur inattendue lors de la recherche: {e}")
            return []

    def extract_result_info(self, result_element: BeautifulSoup) -> Optional[Dict]:
        """Extrait les informations d'un élément de résultat"""
        try:
            info = {
                'title': '',
                'url': '',
                'description': '',
                'metadata': {}
            }

            # Extraction du titre et URL
            title_selectors = [
                'h3 a', 'h2 a', 'h4 a', '.title a',
                'a[href*="/en/"]', '.field--name-title a'
            ]

            title_link = None
            for selector in title_selectors:
                title_link = result_element.select_one(selector)
                if title_link and title_link.get_text(strip=True):
                    break

            if title_link:
                info['title'] = title_link.get_text(strip=True)
                href = title_link.get('href', '')
                if href:
                    info['url'] = urljoin(BASE_URL, href) if not href.startswith('http') else href

            # Extraction de la description
            desc_selectors = [
                '.field--name-body', '.description', '.summary',
                '.field--name-field-meta-description', 'p'
            ]

            for selector in desc_selectors:
                desc_elem = result_element.select_one(selector)
                if desc_elem:
                    info['description'] = desc_elem.get_text(strip=True)[:200]
                    break

            # Extraction des métadonnées (dates, types, etc.)
            metadata_patterns = {
                'date': r'\b\d{1,2}[\/\-]\d{1,2}[\/\-]\d{4}\b|\b\d{4}[\/\-]\d{1,2}[\/\-]\d{1,2}\b',
                'reference': r'\b[A-Z]{2,}[\/\-]\d+[\/\-]\d+\b|\bEMA\/\d+\/\d+\b',
                'type': r'\b(guidance|guideline|procedure|directive|regulation)\b'
            }

            text_content = result_element.get_text()
            for key, pattern in metadata_patterns.items():
                matches = re.findall(pattern, text_content, re.IGNORECASE)
                if matches:
                    info['metadata'][key] = matches[0]

            return info if info['title'] else None

        except Exception as e:
            logger.warning(f"⚠️ Erreur extraction info résultat: {e}")
            return None

    def get_document_details(self, document_url: str) -> Dict:
        """Récupère les détails complets d'un document EMA"""
        try:
            logger.debug(f"📄 Récupération détails: {document_url}")

            response = self.session.get(document_url, timeout=30)
            response.raise_for_status()

            soup = BeautifulSoup(response.text, 'html.parser')

            details = {
                'original_publication_date': None,
                'ema_publication_date': None,
                'ema_reference': '',
                'ema_source_url': document_url,
                'ema_title': '',
                'additional_info': {}
            }

            # Extraction du titre de la page
            title_selectors = ['h1', '.page-title', '.field--name-title', 'title']
            for selector in title_selectors:
                title_elem = soup.select_one(selector)
                if title_elem:
                    details['ema_title'] = title_elem.get_text(strip=True)
                    break

            # Extraction des dates avec différents sélecteurs
            date_selectors = [
                ('.field--name-field-first-published', 'original_publication_date'),
                ('.field--name-field-date', 'original_publication_date'),
                ('.field--name-field-last-updated', 'ema_publication_date'),
                ('.field--name-field-updated', 'ema_publication_date'),
                ('.dates-metadata .first-published', 'original_publication_date'),
                ('.dates-metadata .last-updated', 'ema_publication_date')
            ]

            for selector, field_name in date_selectors:
                date_elem = soup.select_one(selector)
                if date_elem:
                    # Chercher un élément time
                    time_elem = date_elem.select_one('time')
                    if time_elem:
                        date_value = time_elem.get('datetime') or time_elem.get_text(strip=True)
                    else:
                        date_value = date_elem.get_text(strip=True)

                    parsed_date = self.parse_date(date_value)
                    if parsed_date:
                        details[field_name] = parsed_date.isoformat()

            # Extraction de la référence EMA
            ref_selectors = [
                '.field--name-field-ema-reference-number',
                '.field--name-field-reference',
                '[class*="reference"]'
            ]

            for selector in ref_selectors:
                ref_elem = soup.select_one(selector)
                if ref_elem:
                    details['ema_reference'] = ref_elem.get_text(strip=True)
                    break

            # Recherche de référence dans le texte si pas trouvée
            if not details['ema_reference']:
                text_content = soup.get_text()
                ref_patterns = [
                    r'EMA[\/\-]\d+[\/\-]\d+',
                    r'EMEA[\/\-]\d+[\/\-]\d+',
                    r'\b[A-Z]{2,}[\/\-]\d+[\/\-]\d+\b'
                ]

                for pattern in ref_patterns:
                    match = re.search(pattern, text_content)
                    if match:
                        details['ema_reference'] = match.group()
                        break

            # Informations additionnelles
            additional_selectors = {
                'document_type': '.field--name-field-document-type',
                'status': '.field--name-field-status',
                'language': '.field--name-field-language'
            }

            for key, selector in additional_selectors.items():
                elem = soup.select_one(selector)
                if elem:
                    details['additional_info'][key] = elem.get_text(strip=True)

            logger.debug(f"✅ Détails extraits: {details['ema_title'][:50]}...")
            return details

        except Exception as e:
            logger.error(f"❌ Erreur récupération détails: {e}")
            return {
                'original_publication_date': None,
                'ema_publication_date': None,
                'ema_reference': '',
                'ema_source_url': document_url,
                'ema_title': '',
                'additional_info': {}
            }

    def parse_date(self, date_str: str) -> Optional[datetime]:
        """Parse une date dans différents formats"""
        if not date_str:
            return None

        date_str = date_str.strip()

        # Formats de date courants
        date_formats = [
            '%Y-%m-%dT%H:%M:%SZ',  # ISO avec timezone
            '%Y-%m-%dT%H:%M:%S',  # ISO sans timezone
            '%Y-%m-%d',  # ISO simple
            '%d/%m/%Y',  # DD/MM/YYYY
            '%d-%m-%Y',  # DD-MM-YYYY
            '%d %B %Y',  # DD Month YYYY
            '%d %b %Y',  # DD Mon YYYY
            '%B %d, %Y',  # Month DD, YYYY
            '%b %d, %Y'  # Mon DD, YYYY
        ]

        for fmt in date_formats:
            try:
                return datetime.strptime(date_str, fmt)
            except ValueError:
                continue

        # Tentative d'extraction de date avec regex
        date_match = re.search(r'(\d{1,2})[\/\-](\d{1,2})[\/\-](\d{4})', date_str)
        if date_match:
            day, month, year = map(int, date_match.groups())
            try:
                return datetime(year, month, day)
            except ValueError:
                pass

        return None


from urllib.parse import quote_plus
from datetime import datetime
import locale


# Pour parser des dates en français et anglais (EMA souvent en anglais)
def parse_ema_date(date_str: str) -> datetime:
    """Parse une date EMA dans différents formats"""
    if not date_str:
        return None

    date_str = date_str.strip()

    # Formats de date courants sur EMA
    date_formats = [
        '%d/%m/%Y',  # DD/MM/YYYY
        '%d-%m-%Y',  # DD-MM-YYYY
        '%Y-%m-%d',  # YYYY-MM-DD
        '%d %B %Y',  # DD Month YYYY
        '%d %b %Y',  # DD Mon YYYY
        '%B %d, %Y',  # Month DD, YYYY
        '%b %d, %Y'  # Mon DD, YYYY
    ]

    for fmt in date_formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue

    # Tentative d'extraction avec regex pour DD/MM/YYYY
    date_match = re.search(r'(\d{1,2})[\/\-](\d{1,2})[\/\-](\d{4})', date_str)
    if date_match:
        day, month, year = map(int, date_match.groups())
        try:
            return datetime(year, month, day)
        except ValueError:
            pass

    logger.warning(f"⚠️ Impossible de parser la date EMA: '{date_str}'")
    return None


def _empty_ema_result(title: str, search_successful=False, error="", search_url=""):
    """Résultat vide standard pour EMA"""
    return {
        'ema_title': '',
        'ema_source_url': '',
        'search_page_url': search_url,
        'original_publication_date': None,
        'ema_publication_date': None,
        'ema_reference': '',
        'similarity_score': 0.0,
        'search_successful': search_successful,
        'search_performed': True,
        'search_term_used': title,
        'search_timestamp': datetime.now().isoformat(),
        'search_error': error
    }

def fetch_ema_data_from_url(ema_page_url: str) -> dict:
    """
    Extrait les métadonnées directement depuis une page EMA spécifique.
    """
    try:
        logger.info(f"🔍 Extraction directe depuis URL EMA: '{ema_page_url}'")

        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.9',
        }

        response = requests.get(ema_page_url, headers=headers, timeout=30)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")

        # Extraire le titre de la page
        title = ""
        title_selectors = ['h1', '.page-title', '.field--name-title', 'title']
        for selector in title_selectors:
            title_elem = soup.select_one(selector)
            if title_elem:
                title = title_elem.get_text(strip=True)
                if title:
                    break

        # Extraire les dates
        publication_date = None
        last_updated = None

        # Chercher les dates dans la structure de la page
        date_selectors = [
            ('.field--name-field-first-published time', 'first'),
            ('.field--name-field-date time', 'first'),
            ('.field--name-field-last-updated time', 'last'),
            ('.field--name-field-updated time', 'last'),
            ('.dates-metadata .first-published time', 'first'),
            ('.dates-metadata .last-updated time', 'last')
        ]

        for selector, date_type in date_selectors:
            date_elem = soup.select_one(selector)
            if date_elem:
                date_value = date_elem.get('datetime') or date_elem.get_text(strip=True)
                parsed_date = parse_ema_date(date_value)
                if parsed_date:
                    if date_type == 'first' and not publication_date:
                        publication_date = parsed_date
                        logger.debug(f"📅 Date publication trouvée: {parsed_date}")
                    elif date_type == 'last' and not last_updated:
                        last_updated = parsed_date
                        logger.debug(f"📅 Date mise à jour trouvée: {parsed_date}")

        # Extraire la référence EMA
        reference = ''
        reference_selectors = [
            '.field--name-field-ema-reference-number',
            '.field--name-field-reference',
            '[class*="reference"]'
        ]

        for selector in reference_selectors:
            ref_elem = soup.select_one(selector)
            if ref_elem:
                reference = ref_elem.get_text(strip=True)
                if reference:
                    logger.debug(f"📋 Référence trouvée: {reference}")
                    break

        # Chercher le lien PDF sur la page
        pdf_url = ""
        pdf_links = soup.find_all('a', href=lambda x: x and x.endswith('.pdf'))
        if pdf_links:
            pdf_url = pdf_links[0].get('href')
            if pdf_url and not pdf_url.startswith('http'):
                pdf_url = f"https://www.ema.europa.eu{pdf_url}" if pdf_url.startswith('/') else f"https://www.ema.europa.eu/{pdf_url}"

        # Construire le résultat
        result_data = {
            'ema_title': title,
            'ema_source_url': pdf_url or ema_page_url,
            'search_page_url': ema_page_url,
            'original_publication_date': publication_date.isoformat() if publication_date else None,
            'ema_publication_date': last_updated.isoformat() if last_updated else (
                publication_date.isoformat() if publication_date else None),
            'ema_reference': reference,
            'similarity_score': 1.0,  # Score max car on a l'URL directe
            'search_successful': True,
            'search_performed': False,  # Pas de recherche, extraction directe
            'search_term_used': '',
            'search_timestamp': datetime.now().isoformat(),
        }

        logger.info(f"✅ Extraction EMA réussie depuis URL directe")
        return result_data

    except Exception as e:
        logger.error(f"❌ Erreur extraction EMA depuis URL: {e}")
        return {
            'ema_title': '',
            'ema_source_url': '',
            'search_page_url': ema_page_url,
            'original_publication_date': None,
            'ema_publication_date': None,
            'ema_reference': '',
            'similarity_score': 0.0,
            'search_successful': False,
            'search_performed': False,
            'search_term_used': '',
            'search_timestamp': datetime.now().isoformat(),
            'search_error': str(e)
        }



def fetch_ema_dates_by_title(title: str) -> dict:
    """
    Recherche réelle sur le site EMA à partir du titre exact.
    Construit toujours un lien de recherche vers le site officiel EMA.
    """
    search_page_url = ""  # Initialiser en début pour éviter les erreurs

    try:
        logger.info(f"🔍 Recherche EMA pour le titre: '{title}'")

        base_url = "https://www.ema.europa.eu"
        encoded_title = quote_plus(title)
        search_page_url = (
            f"{base_url}/en/search?f%5B0%5D=ema_search_entity_is_document%3ADocument"
            f"&search_api_fulltext={encoded_title}"
        )

        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.9',
            'Accept-Encoding': 'gzip, deflate, br',
            'Connection': 'keep-alive'
        }

        logger.debug(f"🌐 URL de recherche: {search_page_url}")

        # Lancer la recherche HTML
        response = requests.get(search_page_url, headers=headers, timeout=30)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")

        # Chercher le premier résultat avec différents sélecteurs
        result = None
        result_selectors = [
            ".search-results .search-result",
            ".views-row",
            ".item-list .item",
            "[class*='search-result']",
            ".bcl-listing .row > div"
        ]

        for selector in result_selectors:
            result = soup.select_one(selector)
            if result:
                logger.debug(f"✅ Résultat trouvé avec sélecteur: {selector}")
                break

        if not result:
            logger.warning("❌ Aucun résultat EMA trouvé")
            return _empty_ema_result(title, search_successful=False, error="no_result", search_url=search_page_url)

        # Extraire le lien et le titre
        link = result.select_one("a")
        if not link or not link.has_attr('href'):
            logger.warning("❌ Aucun lien trouvé dans le résultat")
            return _empty_ema_result(title, search_successful=False, error="no_link", search_url=search_page_url)

        result_url = base_url + link['href'] if link['href'].startswith('/') else link['href']
        result_title = link.get_text(strip=True) if link else ''

        logger.info(f"📄 Résultat trouvé: '{result_title}' - {result_url}")

        # Extraire la date de publication depuis la page de résultats
        publication_date = None
        date_selectors = [
            ".date-display-single",
            ".field--name-field-date",
            "[class*='date']",
            "time"
        ]

        for selector in date_selectors:
            date_info = result.select_one(selector)
            if date_info:
                date_text = date_info.get_text(strip=True)
                if date_text:
                    publication_date = parse_ema_date(date_text)
                    if publication_date:
                        logger.debug(f"📅 Date trouvée: {date_text} -> {publication_date}")
                        break

        # Aller sur la page du document pour extraire plus de métadonnées
        reference = ''
        last_updated = None

        try:
            logger.debug(f"🔍 Extraction métadonnées depuis: {result_url}")
            doc_page = requests.get(result_url, headers=headers, timeout=15)
            doc_page.raise_for_status()
            doc_soup = BeautifulSoup(doc_page.text, "html.parser")

            # Extraire la référence EMA
            reference_selectors = [
                "div.field--name-field-ema-reference-number",
                ".field--name-field-reference",
                "[class*='reference']"
            ]

            for selector in reference_selectors:
                reference_el = doc_soup.select_one(selector)
                if reference_el:
                    reference = reference_el.get_text(strip=True)
                    if reference:
                        logger.debug(f"📋 Référence trouvée: {reference}")
                        break

            # Extraire la dernière date de mise à jour
            update_selectors = [
                "span.date-display-single",
                ".field--name-field-last-updated time",
                ".field--name-field-updated time",
                "[class*='last-updated'] time"
            ]

            for selector in update_selectors:
                updated_el = doc_soup.select_one(selector)
                if updated_el:
                    date_text = updated_el.get_text(strip=True)
                    if date_text:
                        last_updated = parse_ema_date(date_text)
                        if last_updated:
                            logger.debug(f"🔄 Dernière MAJ: {date_text} -> {last_updated}")
                            break

        except Exception as e:
            logger.warning(f"⚠️ Impossible d'extraire les métadonnées depuis la page EMA: {e}")

        # Construire le résultat final
        result_data = {
            'ema_title': result_title,
            'ema_source_url': result_url,
            'search_page_url': search_page_url,
            'original_publication_date': publication_date.isoformat() if publication_date else None,
            'ema_publication_date': last_updated.isoformat() if last_updated else (
                publication_date.isoformat() if publication_date else None),
            'ema_reference': reference,
            'similarity_score': 1.0,  # Score maximal car c'est le premier résultat
            'search_successful': True,
            'search_performed': True,
            'search_term_used': title,
            'search_timestamp': datetime.now().isoformat(),
        }

        logger.info(f"✅ Recherche EMA réussie pour '{title}'")
        return result_data

    except requests.exceptions.Timeout:
        logger.error(f"⏰ Timeout lors de la recherche EMA pour '{title}'")
        return _empty_ema_result(title, search_successful=False, error="timeout", search_url=search_page_url)

    except requests.exceptions.RequestException as e:
        logger.error(f"🌐 Erreur réseau EMA pour '{title}': {e}")
        return _empty_ema_result(title, search_successful=False, error=f"network_error: {str(e)}",
                                 search_url=search_page_url)

    except Exception as e:
        logger.error(f"❌ Erreur EMA pour '{title}': {e}")
        return _empty_ema_result(title, search_successful=False, error=str(e), search_url=search_page_url)


def _empty_ema_result(title: str, search_successful=False, error="", search_url=""):
    """Résultat vide standard"""
    return {
        'ema_title': '',
        'ema_source_url': '',
        'search_page_url': search_url,
        'original_publication_date': None,
        'ema_publication_date': None,
        'ema_reference': '',
        'similarity_score': 0.0,
        'search_successful': search_successful,
        'search_performed': True,
        'search_term_used': title,
        'search_timestamp': datetime.now().isoformat(),
        'search_error': error
    }


# Test de la fonction
if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)

    test_titles = [
        "Innovation Task Force (ITF) briefing meeting - Instructions on how to apply",
        "Guideline on the requirements for clinical documentation for orally inhaled products",
        "Paxlovid",
        "COVID-19 vaccine guidance"
    ]

    for title in test_titles:
        print(f"\n{'=' * 80}")
        print(f"TEST: '{title}'")
        print('=' * 80)

        result = fetch_ema_dates_by_title(title)

        print(f"✅ Succès: {result['search_successful']}")
        if result['search_successful']:
            print(f"📄 Titre trouvé: '{result['ema_title']}'")
            print(f"🔗 URL: {result['ema_source_url']}")
            print(f"📅 Première publication: {result['original_publication_date']}")
            print(f"📅 Dernière mise à jour: {result['ema_publication_date']}")
            print(f"📋 Référence: {result['ema_reference']}")
            print(f"🎯 Similarité: {result['similarity_score']:.3f}")
            print(f"🔍 Stratégie utilisée: {result.get('strategy_used', 'N/A')}")
        else:
            print(f"❌ Erreur: {result.get('search_error', 'Non spécifiée')}")

        time.sleep(2)  # Pause entre les tests
import time

import os
import re
import json
import time
import logging
import requests

logger = logging.getLogger(__name__)

import logging
import re
import json
from langchain_ollama import ChatOllama

logger = logging.getLogger(__name__)

llm = ChatOllama(base_url="http://localhost:11434", model="mistral", temperature=0.3, top_p=0.9)

def validate_and_fix_confidence_scores(data: dict) -> dict | None:
    if not data or "confidence_scores" not in data:
        return None

    scores = data["confidence_scores"]
    for key, val in scores.items():
        if not isinstance(val, int):
            try:
                scores[key] = int(round(float(val)))
            except Exception:
                logger.warning(f"Score de confiance invalide pour {key}: {val}")
                return None
        if not (0 <= scores[key] <= 100):
            logger.warning(f"Score de confiance hors limite pour {key}: {scores[key]}")
            return None
    return data


def call_mistral_with_confidence(text_chunk: str, document_url: str = "") -> dict | None:
    prompt = f"""
Vous êtes un expert en analyse de documents. Analysez ce document et extrayez les métadonnées avec des scores de confiance.

TEXTE DU DOCUMENT (premiers 2000 chars):
{text_chunk[:2000]}

URL SOURCE: {document_url}

TÂCHE: Retournez UNIQUEMENT un objet JSON avec les métadonnées extraites ET vos scores de confiance:

{{
    "metadata": {{
        "title": "le titre RÉEL complet du document tel qu'il apparaît",
        "document_type": "guideline|regulation|directive|report|procedure|standard|other",
        "publication_date": "date exacte (format DD Month YYYY)",
        "version": "numéro de version/référence du document",
        "source": "EMA pour docs européens, FDA pour docs US, ou organisation réelle",
        "context": "domaine principal (pharmaceutical, medical, legal, etc.)",
        "country": "code pays (EU, US, etc.)",
        "language": "code langue (en, fr, etc.)"
    }},
    "confidence_scores": {{
        "title": 95,
        "document_type": 90,
        "publication_date": 95,
        "version": 20,
        "source": 90,
        "context": 80,
        "country": 95,
        "language": 98
    }},
    "extraction_reasoning": {{
        "title": "Titre trouvé clairement dans l'en-tête du document",
        "document_type": "Document mentionne explicitement 'procedure' dans le titre",
        "publication_date": "Date '6 June 2006' trouvée dans le texte",
        "version": "Numéro de version Rev.11 trouvé",
        "source": "URL indique source européenne, utilisation EMA",
        "context": "Plusieurs termes pharmaceutiques détectés",
        "country": "Domaine URL .eu indique Union Européenne",
        "language": "Texte clairement en anglais"
    }}
}}

RÈGLES STRICTES:
- Le titre doit être EXACTEMENT celui qui apparaît dans le document, complet et précis
- OBLIGATOIRE: Tous les scores doivent être des entiers entre 0 et 100 (inclus)
- 0 = information complètement absente
- 1-30 = très incertain, probablement incorrect
- 31-60 = incertain, nécessite vérification
- 61-80 = assez confiant
- 81-95 = très confiant
- 96-100 = absolument certain
- INTERDIT: aucun score > 100 ou < 0
- INTERDIT: scores décimaux (utilisez seulement des entiers)

Retournez UNIQUEMENT le JSON, sans texte additionnel.
"""

    try:
        response = llm.invoke([{"role": "user", "content": prompt}])
        print(f"Réponse brute Mistral: {response.content[:500]}...")

        # Récupération correcte du contenu
        content = response.content

        match = re.search(r'\{.*\}', content, re.DOTALL)
        if not match:
            logger.warning("❌ Pas de JSON trouvé dans la réponse Mistral")
            return None

        try:
            parsed = json.loads(match.group())

            # LOG DÉTAILLÉ pour debug
            print(f"✅ JSON parsé avec succès")
            print(f"Titre extrait: '{parsed.get('metadata', {}).get('title', 'NON TROUVÉ')}'")
            print(f"Metadata complète: {parsed.get('metadata', {})}")

        except json.JSONDecodeError as e:
            logger.warning(f"❌ Erreur parsing JSON Mistral: {e}")
            return None

        validated_data = validate_and_fix_confidence_scores(parsed)
        if validated_data:
            # VÉRIFICATION FINALE du titre
            final_title = validated_data.get('metadata', {}).get('title', '').strip()
            print(f"🔍 Titre final après validation: '{final_title}'")

            if not final_title or final_title.lower() in ['document sans titre', 'untitled']:
                logger.warning("⚠️ Titre final invalide après validation")
            else:
                logger.info(f"✅ Titre valide confirmé: '{final_title}'")

            logger.info("✅ Réponse Mistral valide reçue et scores validés")
            return validated_data
        else:
            logger.warning("⚠️ Validation des scores échouée")
            return None

    except Exception as e:
        logger.error(f"❌ Erreur inattendue Mistral: {e}")
        return None



# def validate_and_fix_confidence_scores(data):
#     """
#     Valide et corrige les scores de confiance pour s'assurer qu'ils sont entre 0-100.
#     Retourne None si la structure JSON est invalide.
#     """
#     try:
#         if not isinstance(data, dict) or "confidence_scores" not in data:
#             logger.error("❌ Structure JSON invalide: 'confidence_scores' manquant")
#             return None
#
#         confidence_scores = data["confidence_scores"]
#         if not isinstance(confidence_scores, dict):
#             logger.error("❌ 'confidence_scores' doit être un objet")
#             return None
#
#         # Validation et correction des scores
#         fixed_scores = {}
#         for key, score in confidence_scores.items():
#             try:
#                 # Conversion en entier si c'est un float
#                 if isinstance(score, float):
#                     score = int(round(score))
#                 elif not isinstance(score, int):
#                     # Si ce n'est ni int ni float, essayer de convertir
#                     score = int(float(str(score)))
#
#                 # Clamp entre 0 et 100
#                 if score < 0:
#                     logger.warning(f"⚠️ Score négatif pour '{key}': {score} → 0")
#                     score = 0
#                 elif score > 100:
#                     logger.warning(f"⚠️ Score > 100 pour '{key}': {score} → 100")
#                     score = 100
#
#                 fixed_scores[key] = score
#
#             except (ValueError, TypeError) as e:
#                 logger.error(f"❌ Score invalide pour '{key}': {score} - {e}")
#                 # Assigner un score par défaut de 50 (incertain)
#                 fixed_scores[key] = 50
#
#         # Remplacement des scores corrigés
#         data["confidence_scores"] = fixed_scores
#
#         # Vérification que tous les champs requis sont présents
#         required_fields = ["title", "document_type", "publication_date", "version",
#                            "source", "context", "country", "language"]
#
#         missing_scores = [field for field in required_fields if field not in fixed_scores]
#         if missing_scores:
#             logger.warning(f"⚠️ Scores manquants pour: {missing_scores}")
#             # Ajouter des scores par défaut pour les champs manquants
#             for field in missing_scores:
#                 fixed_scores[field] = 50
#
#         logger.info(f"✅ Scores validés: {fixed_scores}")
#         return data
#
#     except Exception as e:
#         logger.error(f"❌ Erreur lors de la validation des scores: {e}")
#         return None
def calculate_overall_quality(confidence_scores):
    """
    Calcule la qualité globale d'extraction à partir des scores de confiance LLM
    """
    if not confidence_scores:
        return 0

    weights = {
        'title': 1.5,
        'document_type': 1.2,  # Changé de 'type' à 'document_type'
        'publication_date': 1.3,
        'source': 1.2,
        'context': 1.0,
        'language': 0.8,
        'country': 0.8,
        'version': 0.7
    }

    total_weighted_score = 0
    total_weight = 0
    for field, score in confidence_scores.items():
        w = weights.get(field, 1.0)
        total_weighted_score += score * w
        total_weight += w

    return int(total_weighted_score / total_weight) if total_weight else 0


def extract_full_text(file_path: str) -> str:
    """Lit tout le PDF, nettoie et enlève les stopwords."""
    try:
        reader = PdfReader(file_path)
        pages = [p.extract_text() or "" for p in reader.pages]
        text = " ".join(pages)
        text = re.sub(r"[^0-9A-Za-zÀ-ÖØ-öø-ÿ\s\.,;:\-'\(\)]", " ", text)
        text = re.sub(r"\s+", " ", text).strip()
        words = text.split()
        return " ".join(w for w in words if w.lower() not in STOPWORDS)
    except Exception as e:
        logger.error(f"Erreur extraction texte: {e}")
        return ""


def extract_metadonnees(file_path: str, source_url: str) -> dict:
    """
    Fonction principale d'extraction avec fallback robuste amélioré
    """
    logger.info(f"🚀 Début extraction pour: {file_path}")

    # Extraction du texte
    full_text = extract_full_text(file_path)
    if not full_text:
        logger.error("❌ Impossible d'extraire le texte")
        return extract_basic_fallback(file_path, source_url)

    # Tentative d'extraction avec Mistral
    llm_result = call_mistral_with_confidence(full_text, source_url)

    if llm_result and 'metadata' in llm_result and 'confidence_scores' in llm_result:
        logger.info("✅ Extraction Mistral réussie")
        metadata = llm_result['metadata']
        conf_scores = llm_result['confidence_scores']
        reasoning = llm_result.get('extraction_reasoning', {})

        metadata['url_source'] = source_url
        metadata['quality'] = {
            'extraction_rate': calculate_overall_quality(conf_scores),
            'field_scores': conf_scores,
            'extraction_reasoning': reasoning,
            'extracted_fields': sum(1 for s in conf_scores.values() if s >= 50),
            'total_fields': len(conf_scores),
            'llm_powered': True
        }
        return metadata
    else:
        # Fallback vers extraction robuste par patterns
        logger.warning("⚠️ Fallback vers extraction robuste")
        return extract_robust_fallback(file_path, source_url, full_text)


def extract_robust_fallback(file_path: str, source_url: str, full_text: str = None) -> dict:
    """Fallback robuste avec extraction par patterns"""
    logger.info("🔧 Utilisation extraction robuste par patterns")

    if not full_text:
        full_text = extract_full_text(file_path)

    # Obtenir le titre du document depuis PDF metadata
    try:
        reader = PdfReader(file_path)
        document_title = reader.metadata.title if reader.metadata else None
    except:
        document_title = Path(file_path).stem

    # Utiliser l'extracteur robuste
    extractor = RobustMetadataExtractor()
    metadata = extractor.extract_robust_patterns(full_text, "pdf", document_title)

    # Ajouter les champs manquants
    metadata['url_source'] = source_url
    metadata['quality'] = {
        'extraction_rate': calculate_overall_quality(metadata['confidence_scores']),
        'field_scores': metadata['confidence_scores'],
        'extraction_reasoning': {
            'method': 'Extraction robuste par patterns',
            'fallback_reason': 'LLM non disponible ou échec'
        },
        'extracted_fields': sum(1 for s in metadata['confidence_scores'].values() if s >= 50),
        'total_fields': len(metadata['confidence_scores']),
        'llm_powered': False
    }

    return metadata


def extract_basic_fallback(file_path: str, source_url: str) -> dict:
    """Fallback basique avec scores de confiance honnêtes"""
    logger.warning("⚠️ Utilisation fallback basique")

    try:
        reader = PdfReader(file_path)
        info = reader.metadata or {}
        title = info.title or Path(file_path).stem

        # Essayer de détecter la langue sur un échantillon
        sample_text = ""
        try:
            for i, page in enumerate(reader.pages[:3]):  # 3 premières pages
                sample_text += page.extract_text() or ""
            lang = detect(sample_text) if sample_text else "en"
        except:
            lang = "en"
    except:
        title = Path(file_path).stem
        lang = "en"

    basic_meta = {
        "title": title,
        "document_type": "other",  # Changé de "type" à "document_type"
        "publication_date": "",
        "version": "",
        "source": "",
        "context": "general",
        "country": "",
        "language": lang,
        "url_source": source_url
    }

    conf_scores = {
        'title': 30 if title and title != Path(file_path).stem else 10,
        'document_type': 0,  # Changé de 'type' à 'document_type'
        'publication_date': 0,
        'version': 0,
        'source': 0,
        'context': 0,
        'country': 0,
        'language': 70 if lang else 0
    }

    basic_meta['quality'] = {
        'extraction_rate': calculate_overall_quality(conf_scores),
        'field_scores': conf_scores,
        'extraction_reasoning': {
            'title': 'Titre basique des métadonnées PDF',
            'method': 'Extraction basique de fallback',
            'limitation': 'Impossible d\'analyser le contenu'
        },
        'extracted_fields': 1 if title else 0,
        'total_fields': len(conf_scores),
        'llm_powered': False
    }

    return basic_meta


# Fonction de test
if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)

    # Test avec un fichier exemple
    test_file = "example.pdf"
    test_url = "https://example.com/doc.pdf"

    if os.path.exists(test_file):
        result = extract_metadonnees(test_file, test_url)
        print("Résultat extraction:")
        print(json.dumps(result, indent=2, ensure_ascii=False))
    else:
        print("Fichier de test non trouvé")